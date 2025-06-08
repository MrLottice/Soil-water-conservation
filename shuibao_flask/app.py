from flask import Flask, render_template, request, jsonify, send_file, Response
import requests
import json
import os
from datetime import datetime
from docx import Document
from docx.shared import Pt
import markdown
import io
import subprocess
import webbrowser
from flask_cors import CORS  # 导入CORS扩展


# 创建Flask应用实例
app = Flask(__name__)
# 添加CORS支持，允许所有来源
CORS(app, resources={r"/*": {"origins": "*"}})

# 全局变量存储当前文档
current_doc = None
current_filename = None

# Dify API配置
DIFY_API_URL = "http://localhost/v1/workflows/run"  # 工作流API地址
DIFY_API_KEY_SHUIBAO = "app-NDoLGNo1764JD79Vv6WRuNhl" # 水保方案
DIFY_UPLOAD_URL = "http://localhost/v1/files/upload"  # Dify文件上传API地址


# 路由：首页
@app.route('/test', methods=['GET'])
def home():
    return 'test'

@app.route('/shuibao', methods=['GET', 'POST'])
def shuibao():
    global current_doc, current_filename
    
    raw_text = request.get_data(as_text=True)
    print("\n=== 开始处理请求 ===")
    print("收到文本内容：")
    print(raw_text)
    
    try:
        # 尝试解析JSON数据
        try:
            data = json.loads(raw_text)
            content = data.get('content', '')
            is_final = data.get('is_final', False)
            print(f"\n解析JSON成功:")
            print(f"- is_final: {is_final}")
            print(f"- content长度: {len(content)}")
        except json.JSONDecodeError as e:
            print(f"\nJSON解析错误: {str(e)}")
            content = raw_text
            is_final = False
            
        if not content:
            print("错误：内容为空")
            return jsonify({
                'message': 'error',
                'status': '内容为空'
            }), 400
        
        # 如果是第一段，创建新文档
        if current_doc is None:
            print("\n创建新文档")
            current_doc = Document()
            # 设置默认字体
            style = current_doc.styles['Normal']
            font = style.font
            font.name = '宋体'
            font.size = Pt(10.5)
            # 生成文件名
            timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
            current_filename = f'水土保持方案_{timestamp}.docx'
            print(f"新文档文件名: {current_filename}")
        else:
            print(f"\n继续使用现有文档: {current_filename}")
        
        # 处理内容并添加到文档
        print("\n开始处理内容")
        lines = content.split('\n')
        print(f"处理行数: {len(lines)}")
        
        for line in lines:
            if line.startswith('###'):
                current_doc.add_heading(line.replace('###', '').strip(), level=3)
            elif line.startswith('##'):
                current_doc.add_heading(line.replace('##', '').strip(), level=2)
            elif line.startswith('#'):
                current_doc.add_heading(line.replace('#', '').strip(), level=1)
            elif line.startswith('|'):
                if '|' in line:
                    table_data = [cell.strip() for cell in line.split('|') if cell.strip()]
                    if len(table_data) > 0:
                        table = current_doc.add_table(rows=1, cols=len(table_data))
                        table.style = 'Table Grid'
                        for i, cell in enumerate(table_data):
                            table.cell(0, i).text = cell
            elif line.strip():
                current_doc.add_paragraph(line.strip())
        
        # 保存到本地
        save_dir = r'D:\水保项目\shuibao_flask'
        if not os.path.exists(save_dir):
            print(f"创建输出目录: {save_dir}")
            os.makedirs(save_dir)
        
        save_path = os.path.join(save_dir, current_filename)
        print(f"\n准备保存文档到: {save_path}")
        try:
            current_doc.save(save_path)
            print(f"文档已保存到: {save_path}")
            if os.path.exists(save_path):
                file_size = os.path.getsize(save_path)
                print(f"文档保存成功，文件大小: {file_size} 字节")
            else:
                print("错误：文档保存失败，文件不存在")
        except Exception as e:
            print(f"保存文档时出错: {str(e)}")
            return jsonify({
                'message': 'error',
                'status': f'保存文档失败: {str(e)}'
            }), 500
        
        # 如果是最后一段，返回文件并重置文档
        if is_final:
            print("\n=== 文档生成完成 ===")
            print(f"文档路径: {save_path}")
            
            # 确保文档已保存
            if os.path.exists(save_path):
                file_size = os.path.getsize(save_path)
                print(f"文档已成功保存，文件大小: {file_size} 字节")
                
                # 尝试打开文档
                try:
                    abs_path = os.path.abspath(save_path)
                    print(f"正在打开文档: {abs_path}")
                    
                    # 使用Windows的start命令打开文档
                    subprocess.Popen(['start', '', abs_path], shell=True)
                    print("已发送打开文档命令")
                except Exception as e:
                    print(f"打开文档时出错: {str(e)}")
                    print("请手动打开文档:", abs_path)
            else:
                print("错误：文档未成功保存")
                print("当前工作目录:", os.getcwd())
                print("尝试创建的文件路径:", save_path)
                return jsonify({
                    'message': 'error',
                    'status': '文档保存失败'
                }), 500
            
            # 准备返回响应
            response = send_file(
                save_path,
                mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                as_attachment=True,
                download_name=current_filename
            )
            
            # 重置文档状态
            current_doc = None
            current_filename = None
            print("文档状态已重置")
            
            return response
        
        print("\n=== 请求处理完成 ===")
        return jsonify({
            'message': 'success',
            'status': '内容已添加到文档'
        })
        
    except Exception as e:
        print(f"\n发生错误: {str(e)}")
        import traceback
        print("错误详情:")
        print(traceback.format_exc())
        return jsonify({
            'message': 'error',
            'status': str(e)
        }), 500

@app.route('/shuibao_api', methods=['POST'])
def shuibao_api():
    try:
        # 获取请求数据
        if request.is_json:
            data = request.get_json()
        else:
            data = request.form.to_dict()
            
        if not data or 'name' not in data:
            return jsonify({
                'message': 'error',
                'status': '请求数据格式错误，必须包含name字段'
            }), 400

        # 构建符合Dify API规范的请求体
        dify_payload = {
            'inputs': {
                'name': data.get('name', '').strip('"')  # 移除可能的引号
            },
            'response_mode': 'streaming',  # 使用流式模式
            'user': 'user_' + datetime.now().strftime('%Y%m%d%H%M%S')
        }

        print("\n=== 发送请求到Dify水保API ===")
        print(f"请求格式: {'JSON' if request.is_json else 'form-data'}")
        print(f"项目名称: {data.get('name', '')}")
        print(f"请求体: {json.dumps(dify_payload, ensure_ascii=False, indent=2)}")

        # 准备发送到Dify的请求
        dify_headers = {
            'Authorization': f'Bearer {DIFY_API_KEY_SHUIBAO}',
            'Content-Type': 'application/json'
        }

        # 发送请求到Dify并获取流式响应
        dify_response = requests.post(
            DIFY_API_URL,
            headers=dify_headers,
            json=dify_payload,
            stream=True  # 启用流式传输
        )

        # 检查Dify响应
        if dify_response.status_code != 200:
            print(f"Dify API错误: {dify_response.status_code}")
            print(f"错误信息: {dify_response.text}")
            return jsonify({
                'message': 'error',
                'status': f'Dify API错误: {dify_response.status_code}',
                'error_detail': dify_response.text
            }), 500

        def generate():
            try:
                full_text = ""  # 用于存储完整的文本
                print("\n=== 开始生成水保方案文本 ===")
                for line in dify_response.iter_lines():
                    if line:
                        line = line.decode('utf-8')
                        if line.startswith('data: '):
                            try:
                                data = json.loads(line[6:])  # 去掉'data: '前缀
                                if data.get('event') == 'text_chunk' and 'data' in data and 'text' in data['data']:
                                    text_chunk = data['data']['text']
                                    full_text += text_chunk  # 拼接文本
                                    # 打印到控制台
                                    print(text_chunk, end='', flush=True)
                                    # 构造返回给前端的数据
                                    response_data = {
                                        'type': 'text',
                                        'content': text_chunk,
                                        'full_text': full_text
                                    }
                                    yield f"data: {json.dumps(response_data, ensure_ascii=False)}\n\n"
                                elif data.get('event') == 'done':
                                    # 发送完成事件
                                    response_data = {
                                        'type': 'done',
                                        'full_text': full_text
                                    }
                                    yield f"data: {json.dumps(response_data, ensure_ascii=False)}\n\n"
                                else:
                                    # 其他类型的事件直接转发
                                    yield f"data: {json.dumps(data, ensure_ascii=False)}\n\n"
                            except json.JSONDecodeError as e:
                                print(f"\nJSON解析错误: {str(e)}")
                                continue
                print("\n=== 水保方案文本生成完成 ===")
            except Exception as e:
                print(f"\n流式处理错误: {str(e)}")
                error_data = {
                    'type': 'error',
                    'message': str(e)
                }
                yield f"data: {json.dumps(error_data, ensure_ascii=False)}\n\n"

        return Response(generate(), mimetype='text/event-stream')

    except Exception as e:
        print(f"\n发生错误: {str(e)}")
        import traceback
        print("错误详情:")
        print(traceback.format_exc())
        return jsonify({
            'message': 'error',
            'status': str(e)
        }), 500



# 错误处理
@app.errorhandler(404)
def not_found_error(error):
    return jsonify({
        'message': 'Resource not found',
        'status': 'error'
    }), 404

if __name__ == '__main__':
    app.run(debug=True, host='127.0.0.1', port=5000) 