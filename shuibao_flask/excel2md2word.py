import pandas as pd
import os
from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.enum.table import WD_TABLE_ALIGNMENT

def excel_to_word(excel_path, output_docx_path):
    """将Excel转换为Word文档"""
    try:
        # 检查文件是否存在
        if not os.path.exists(excel_path):
            raise FileNotFoundError(f"Excel文件不存在: {excel_path}")

        # 创建Word文档
        doc = Document()
        
        # 设置页面边距（单位：厘米）
        sections = doc.sections
        for section in sections:
            section.top_margin = Cm(2)
            section.bottom_margin = Cm(2)
            section.left_margin = Cm(2)
            section.right_margin = Cm(2)

        # 读取Excel文件
        excel_file = pd.ExcelFile(excel_path)
        
        # 处理每个表单
        for sheet_name in excel_file.sheet_names:
            print(f"正在处理表单: {sheet_name}")
            
            # 读取数据，跳过第一行作为标题
            df = pd.read_excel(excel_path, sheet_name=sheet_name, header=0)
            
            # 处理空值
            df = df.fillna('')
            
            # 添加表单标题
            heading = doc.add_heading(f"表单：{sheet_name}", level=2)
            heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # 创建表格
            table = doc.add_table(rows=len(df)+1, cols=len(df.columns))
            table.style = 'Table Grid'
            
            # 设置表头
            for col in range(len(df.columns)):
                cell = table.cell(0, col)
                # 获取列名，如果是Unnamed则设为空
                col_name = str(df.columns[col])
                if 'Unnamed' in col_name:
                    col_name = ''
                cell.text = col_name
                # 设置表头格式
                paragraph = cell.paragraphs[0]
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = paragraph.runs[0] if paragraph.runs else paragraph.add_run()
                run.font.bold = True
                run.font.size = Pt(11)
                run.font.name = '宋体'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
            
            # 填充数据
            for row in range(len(df)):
                for col in range(len(df.columns)):
                    cell = table.cell(row+1, col)
                    value = df.iloc[row, col]
                    # 处理数值格式
                    if isinstance(value, (int, float)):
                        cell.text = f"{value:,.2f}" if isinstance(value, float) else str(value)
                    else:
                        cell.text = str(value)
                    # 设置单元格格式
                    paragraph = cell.paragraphs[0]
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run = paragraph.runs[0] if paragraph.runs else paragraph.add_run()
                    run.font.size = Pt(10)
                    run.font.name = '宋体'
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
            
            # 添加空行
            doc.add_paragraph()
        
        # 保存文档
        doc.save(output_docx_path)
        print(f"已成功生成Word文档：{output_docx_path}")
        return True
        
    except Exception as e:
        print(f"转换过程中出错: {str(e)}")
        return False

if __name__ == "__main__":
    # 设置输入输出路径
    base_dir = r"D:\水保项目\巴塘水保\水保需材料\水保需材料\（二期）清单 控价\【清单控价送审电子档】（二期）"
    excel_file = os.path.join(base_dir, "5-建设项目预算总投资费用组成表.xls")
    output_docx = os.path.join(base_dir, "5-建设项目预算总投资费用组成表.docx")
    
    try:
        if excel_to_word(excel_file, output_docx):
            print("转换成功完成！")
        else:
            print("转换过程中出现错误。")
    except Exception as e:
        print(f"程序执行失败: {str(e)}")